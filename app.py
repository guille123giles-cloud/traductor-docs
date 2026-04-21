import streamlit as st
import tempfile
import os
import time
from PIL import Image
from deep_translator import GoogleTranslator
import datetime

# --- NUEVA LIBRERÍA DE GOOGLE ---
from google import genai

# Librerías para traducción de documentos sin Azure
import io
import docx
import pdfplumber
from fpdf import FPDF

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Traductor de documentos", page_icon="🌍", layout="wide")

def traducir_bloque(texto, origen, destino):
    """Traduce un bloque de texto entero. Súper rápido y con reintentos."""
    if not texto.strip() or len(texto.strip()) < 2: 
        return texto
    for intento in range(3):
        try:
            return GoogleTranslator(source=origen, target=destino).translate(texto)
        except Exception:
            time.sleep(1.5)
    return texto 

def _traducir_parrafos(parrafos, origen, destino, barra, progreso_inicio, progreso_fin):
    """Traduce una lista de párrafos respetando los vacíos, con barra de progreso."""
    total = len(parrafos)
    traducidos = []
    for i, p in enumerate(parrafos):
        traducidos.append(traducir_bloque(p, origen, destino))
        progreso = progreso_inicio + int((i + 1) / total * (progreso_fin - progreso_inicio))
        barra.progress(progreso)
    return traducidos

def procesar_documento(archivo_subido, origen, destino, barra, estado):
    """Traduce archivos PDF y DOCX de forma 100% gratuita, sin APIs de pago."""
    extension = os.path.splitext(archivo_subido.name)[1].lower()

    try:
        # ── DOCX ────────────────────────────────────────────────────────────────
        if extension == ".docx":
            estado.info("⚙️ Paso 1/3: Leyendo el documento Word...")
            barra.progress(10)

            doc_in = docx.Document(io.BytesIO(archivo_subido.getvalue()))
            parrafos_texto = [p.text for p in doc_in.paragraphs]

            estado.info("🚀 Paso 2/3: Traduciendo párrafos (puede demorar según el tamaño)...")
            traducidos = _traducir_parrafos(parrafos_texto, origen, destino, barra, 10, 85)

            estado.info("💾 Paso 3/3: Reconstruyendo el documento con el mismo formato...")
            barra.progress(90)

            # Reemplazar el texto de cada párrafo conservando el estilo
            for p_obj, texto_nuevo in zip(doc_in.paragraphs, traducidos):
                if not p_obj.runs:
                    continue
                # Limpiar runs intermedios y dejar solo el primero con el texto traducido
                for run in p_obj.runs[1:]:
                    run.text = ""
                p_obj.runs[0].text = texto_nuevo

            # Traducir también las tablas
            for tabla in doc_in.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for p_obj in celda.paragraphs:
                            texto_celda = p_obj.text
                            if texto_celda.strip():
                                traducido = traducir_bloque(texto_celda, origen, destino)
                                if p_obj.runs:
                                    for run in p_obj.runs[1:]:
                                        run.text = ""
                                    p_obj.runs[0].text = traducido

            buffer = io.BytesIO()
            doc_in.save(buffer)
            buffer.seek(0)

            estado.success("✅ ¡Traducción del documento Word completada!")
            barra.progress(100)
            return buffer.getvalue()

        # ── PDF ─────────────────────────────────────────────────────────────────
        elif extension == ".pdf":
            estado.info("⚙️ Paso 1/3: Extrayendo texto del PDF...")
            barra.progress(10)

            paginas_texto = []
            with pdfplumber.open(io.BytesIO(archivo_subido.getvalue())) as pdf:
                for pagina in pdf.pages:
                    texto = pagina.extract_text() or ""
                    paginas_texto.append(texto)

            if not any(p.strip() for p in paginas_texto):
                estado.error("❌ El PDF parece ser una imagen escaneada sin texto extraíble. Convertilo a DOCX o usá la opción de imagen.")
                return None

            estado.info("🚀 Paso 2/3: Traduciendo el contenido del PDF...")
            paginas_traducidas = _traducir_parrafos(paginas_texto, origen, destino, barra, 10, 85)

            estado.info("💾 Paso 3/3: Generando el nuevo PDF traducido...")
            barra.progress(90)

            def _limpiar_linea(texto):
                if not texto:
                    return ""
                return texto.encode("latin-1", errors="replace").decode("latin-1")

            def _escribir_linea(pdf, linea):
                linea = _limpiar_linea(linea)
                if not linea.strip():
                    pdf.ln(4)
                    return
                palabras = linea.split(" ")
                partes = []
                for palabra in palabras:
                    while len(palabra) > 55:
                        partes.append(palabra[:55])
                        palabra = palabra[55:]
                    partes.append(palabra)
                linea_segura = " ".join(partes)
                try:
                    pdf.multi_cell(0, 6, linea_segura)
                except Exception:
                    pass

            pdf_out = FPDF()
            pdf_out.set_margins(20, 20, 20)
            pdf_out.set_auto_page_break(auto=True, margin=20)
            pdf_out.add_page()
            pdf_out.set_font("Helvetica", size=10)

            for i, texto_pag in enumerate(paginas_traducidas):
                if i > 0:
                    pdf_out.add_page()
                for linea in texto_pag.split("\n"):
                    _escribir_linea(pdf_out, linea)

            buffer = io.BytesIO()
            pdf_out.output(buffer)
            buffer.seek(0)

            estado.success("✅ ¡Traducción del PDF completada!")
            barra.progress(100)
            return buffer.getvalue()

    except Exception as e:
        estado.error(f"❌ Ocurrió un error al procesar el documento: {e}")
        return None

def procesar_imagen(imagen_subida, origen, destino, barra, estado):
    """Extrae texto de una imagen usando Google Gemini y lo traduce en bloque."""
    estado.info("⚙️ Paso 1/2: Leyendo la imagen con Inteligencia Artificial...")
    img = Image.open(imagen_subida)
    
    try:
        # Usar el nuevo cliente de Google GenAI
        client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
        prompt_ocr = "Extrae todo el texto de esta imagen con precisión absoluta. Respeta los saltos de línea y el orden de los mensajes. No agregues introducciones, ni comentarios, solo devuelve el texto exacto que ves."
        
        respuesta = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt_ocr, img]
        )
        
        texto_extraido = respuesta.text
        
    except Exception as e:
        return None, f"⚠️ Error al leer la imagen con Gemini: {e}"
    
    if not texto_extraido or not texto_extraido.strip():
        return None, "⚠️ No se detectó texto legible en la captura."

    estado.info("🚀 Paso 2/2: Traduciendo texto de forma optimizada...")
    barra.progress(50)
    
    # TRADUCCIÓN EN UN SOLO VIAJE (Muchísimo más rápido)
    texto_final = traducir_bloque(texto_extraido, origen, destino)
            
    barra.progress(100)
    estado.success("✅ ¡Lectura y traducción de imagen completada al instante!")
    return texto_extraido, texto_final


# --- INTERFAZ PARA EL CLIENTE ---
st.title("🌍 Traductor de Documentos e Imágenes")
st.divider()

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("1. Subir Archivo")
    archivo = st.file_uploader("Soporta: PDF, DOCX, PNG, JPG, JPEG", type=["pdf", "docx", "png", "jpg", "jpeg"])

with col2:
    st.subheader("2. Opciones de Idioma")
    with st.container(border=True):
        idiomas = {"Auto-detectar": "auto", "Español": "es", "Inglés": "en", "Portugués": "pt", "Francés": "fr", "Italiano": "it", "Alemán": "de"}
        ori_nombre = st.selectbox("Idioma de origen:", list(idiomas.keys()))
        des_nombre = st.selectbox("Traducir al:", [k for k in idiomas.keys() if k != "Auto-detectar"], index=0)
        
        idioma_ori = idiomas[ori_nombre]
        idioma_des = idiomas[des_nombre]

st.divider()

# --- PROCESAMIENTO ---
_, col_btn, _ = st.columns([1, 2, 1])

with col_btn:
    if archivo:
        if archivo.size > 5 * 1024 * 1024:
            st.warning("⚠️ El archivo supera los 5MB. El proceso puede demorar más de lo habitual.")
            
        if st.button("🚀 COMENZAR TRADUCCIÓN", use_container_width=True):
            estado = st.empty()
            barra = st.progress(0)
            
            extension = os.path.splitext(archivo.name)[1].lower()
            nombre_base = os.path.splitext(archivo.name)[0]
            
            # --- RUTA 1: DOCUMENTOS (PDF / DOCX) ---
            if extension in [".pdf", ".docx"]:
                resultado = procesar_documento(archivo, idioma_ori, idioma_des, barra, estado)
                
                if resultado:
                    tipo_mime = "application/pdf" if extension == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    
                    st.download_button(
                        label=f"⬇️ DESCARGAR DOCUMENTO TRADUCIDO ({extension.upper()})", 
                        data=resultado, 
                        file_name=f"TRADUCIDO_{nombre_base}{extension}",
                        mime=tipo_mime,
                        use_container_width=True
                    )
            
            # --- RUTA 2: IMÁGENES (PNG / JPG / JPEG) ---
            elif extension in [".png", ".jpg", ".jpeg"]:
                texto_original, texto_traducido = procesar_imagen(archivo, idioma_ori, idioma_des, barra, estado)
                
                if texto_original:
                    st.markdown("### 📝 Resultados")
                    col_res1, col_res2 = st.columns(2)
                    with col_res1:
                        st.text_area("Texto Original", value=texto_original, height=350)
                    with col_res2:
                        st.text_area(f"Traducción al {des_nombre}", value=texto_traducido, height=350)
                else:
                    if texto_traducido:
                        estado.error(texto_traducido)
